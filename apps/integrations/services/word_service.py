import requests
import logging
from django.conf import settings

logger = logging.getLogger(__name__)


class WordIntegrationService:
    """
    Microsoft Word Online integration service
    """
    
    def __init__(self, integration):
        self.integration = integration
        self.access_token = integration.access_token
        self.graph_api_url = "https://graph.microsoft.com/v1.0"
    
    def get_edit_url(self, contract):
        """
        Get URL to open contract in Word Online
        """
        try:
            # Upload to OneDrive if not already there
            file_id = self._upload_to_onedrive(contract)
            
            # Get Word Online edit URL
            url = f"{self.graph_api_url}/me/drive/items/{file_id}"
            headers = {
                'Authorization': f'Bearer {self.access_token}',
            }
            
            response = requests.get(url, headers=headers)
            response.raise_for_status()
            
            data = response.json()
            return data.get('webUrl')
            
        except Exception as e:
            logger.error(f"Error getting Word edit URL: {str(e)}")
            return None
    
    def _upload_to_onedrive(self, contract):
        """
        Upload contract file to OneDrive
        """
        try:
            # Download from S3
            from apps.contracts.services.upload_service import UploadService
            upload_service = UploadService()
            
            # Get file content
            # In production, download from S3
            file_path = contract.file_path
            
            # Upload to OneDrive
            url = f"{self.graph_api_url}/me/drive/root:/ChainSight/{contract.original_filename}:/content"
            headers = {
                'Authorization': f'Bearer {self.access_token}',
                'Content-Type': 'application/octet-stream'
            }
            
            # Mock response for now
            # In production, actually upload file
            return "mock_file_id"
            
        except Exception as e:
            logger.error(f"Error uploading to OneDrive: {str(e)}")
            raise
    
    def save_from_word(self, contract, file_content, track_changes=True):
        """
        Save updated contract from Word Online
        """
        try:
            # Save new version to S3
            from apps.contracts.services.upload_service import UploadService
            upload_service = UploadService()
            
            # Create new version
            # Store with version number
            version = contract.metadata.get('version', 1) + 1
            new_file_path = f"{contract.file_path}_v{version}"
            
            # Update contract
            contract.file_path = new_file_path
            contract.metadata['version'] = version
            contract.metadata['last_edited_in'] = 'microsoft_word'
            contract.save()
            
            # Log integration activity
            from apps.integrations.models import IntegrationLog
            IntegrationLog.objects.create(
                tenant=contract.tenant,
                integration=self.integration,
                action='import',
                status='success',
                request_data={'contract_id': str(contract.id)}
            )
            
            return contract
            
        except Exception as e:
            logger.error(f"Error saving from Word: {str(e)}")
            raise
    
    def export_template_to_word(self, template):
        """
        Export template as Word document
        """
        try:
            from docx import Document
            from io import BytesIO
            
            # Create Word document
            doc = Document()
            doc.add_heading(template.name, 0)
            
            # Add template content
            for paragraph in template.template_text.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            
            # Save to BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Upload to OneDrive
            url = f"{self.graph_api_url}/me/drive/root:/ChainSight/Templates/{template.name}.docx:/content"
            headers = {
                'Authorization': f'Bearer {self.access_token}',
                'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            }
            
            response = requests.put(url, headers=headers, data=buffer.getvalue())
            response.raise_for_status()
            
            return response.json().get('webUrl')
            
        except Exception as e:
            logger.error(f"Error exporting to Word: {str(e)}")
            raise


class GoogleDocsService:
    """
    Google Docs integration service
    """
    
    def __init__(self, integration):
        self.integration = integration
        self.access_token = integration.access_token
    
    def get_edit_url(self, contract):
        """
        Get URL to open contract in Google Docs
        """
        try:
            # Upload to Google Drive
            file_id = self._upload_to_drive(contract)
            
            # Return Google Docs URL
            return f"https://docs.google.com/document/d/{file_id}/edit"
            
        except Exception as e:
            logger.error(f"Error getting Google Docs URL: {str(e)}")
            return None
    
    def _upload_to_drive(self, contract):
        """
        Upload contract to Google Drive
        """
        try:
            # Google Drive API implementation
            # Mock for now
            return "mock_google_doc_id"
            
        except Exception as e:
            logger.error(f"Error uploading to Google Drive: {str(e)}")
            raise
    
    def sync_from_google_docs(self, contract, document_id):
        """
        Sync updated contract from Google Docs
        """
        try:
            # Download from Google Docs
            url = f"https://www.googleapis.com/drive/v3/files/{document_id}/export?mimeType=application/pdf"
            headers = {
                'Authorization': f'Bearer {self.access_token}'
            }
            
            response = requests.get(url, headers=headers)
            response.raise_for_status()
            
            # Save to S3 and update contract
            # Implementation similar to Word service
            
            return contract
            
        except Exception as e:
            logger.error(f"Error syncing from Google Docs: {str(e)}")
            raise

